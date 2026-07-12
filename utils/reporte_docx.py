"""Generador de informe profesional DOCX con diseño institucional."""
import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import streamlit as st


# Constantes de diseño
COLOR_PRINCIPAL = RGBColor(155, 27, 48)     # #9B1B30 — crimson institucional
COLOR_SECUNDARIO = RGBColor(107, 144, 128)   # #6B9080 — sage
COLOR_TEXTO = RGBColor(11, 15, 25)           # #0B0F19 — deep
COLOR_GRIS = RGBColor(160, 144, 136)         # #A09088 — stone
COLOR_FONDO_TABLA = RGBColor(240, 234, 228)  # #F0EAE4 — cream
FONT_TITULOS = 'Calibri'
FONT_CUERPO = 'Calibri'
MARGEN = Cm(2.5)


def _aplicar_estilo_base(doc: Document):
    """Configura márgenes y estilos base del documento."""
    for section in doc.sections:
        section.top_margin = MARGEN
        section.bottom_margin = MARGEN
        section.left_margin = MARGEN
        section.right_margin = MARGEN
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # Estilo Normal
    style = doc.styles['Normal']
    style.font.name = FONT_CUERPO
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXTO
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Estilo Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT_TITULOS
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRINCIPAL
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    # Estilo Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = FONT_TITULOS
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_PRINCIPAL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)

    # Estilo Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = FONT_TITULOS
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_SECUNDARIO
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)


def _agregar_marca_agua(doc: Document, texto: str = "ESTUDIO CONTEXTO"):
    """Agrega marca de agua diagonal en cada pagina usando XML OOXML."""
    # Registrar namespaces necesarios para VML
    _vml_ns = " ".join(f'xmlns:{pfx}="{uri}"' for pfx, uri in {
        'v': 'urn:schemas-microsoft-com:vml',
        'o': 'urn:schemas-microsoft-com:office:office',
    }.items())

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Watermark como texto rotado en el header
        xml = (
            f'<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' {_vml_ns}>'
            f'  <v:shape id="watermark" style="position:absolute;margin-left:0;margin-top:0;'
            f'    width:450pt;height:150pt;rotation:315;z-index:-251658240;'
            f'    mso-position-horizontal:center;mso-position-vertical:center" '
            f'    fillcolor="#A09088" stroked="f" opacity=".08">'
            f'    <v:textbox inset="0,0,0,0">'
            f'      <w:txbxContent>'
            f'        <w:p>'
            f'          <w:pPr><w:jc w:val="center"/></w:pPr>'
            f'          <w:r>'
            f'            <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
            f'            <w:sz w:val="72"/><w:color w:val="A09088"/></w:rPr>'
            f'            <w:t xml:space="preserve">{texto}</w:t>'
            f'          </w:r>'
            f'        </w:p>'
            f'      </w:txbxContent>'
            f'    </v:textbox>'
            f'  </v:shape>'
            f'</w:pict>'
        )
        run = p.add_run()
        run._r.append(parse_xml(xml))


def _agregar_encabezado(doc: Document, titulo: str = "INFORME DE PERTINENCIA EDUCATIVA"):
    """Agrega encabezado con linea y titulo."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("ESTUDIO CONTEXTO")
        run.font.name = FONT_TITULOS
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GRIS
        run.font.bold = True


def _agregar_pie_pagina(doc: Document):
    """Agrega pie de pagina con numero de pagina."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Linea separadora
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:space="4" w:color="A09088"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        run = p.add_run("— ")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GRIS

        # Numero de pagina
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1 = p.add_run()
        run1._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2 = p.add_run()
        run2._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3 = p.add_run()
        run3._r.append(fldChar2)

        run4 = p.add_run(" —")
        run4.font.size = Pt(8)
        run4.font.color.rgb = COLOR_GRIS


def _portada(doc: Document, nbc: str, depto: str, fecha: str):
    """Genera portada institucional."""
    # Espaciado superior
    for _ in range(6):
        doc.add_paragraph()

    # Linea decorativa
    p_linea = doc.add_paragraph()
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_linea.add_run("_" * 60)
    run.font.color.rgb = COLOR_PRINCIPAL
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Titulo principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run("INFORME DE\nPERTINENCIA EDUCATIVA")
    run.font.name = FONT_TITULOS
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRINCIPAL

    doc.add_paragraph()

    # Subtitulo
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Sistema de Analisis para Estudio de Contexto")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_SECUNDARIO

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata
    metadata = [
        ("NBC Analizado", nbc or "No especificado"),
        ("Cobertura Geografica", depto or "Nacional"),
        ("Fecha de Generacion", fecha),
        ("Version", "2.0"),
        ("Metodologia", "SNIES, SIET, CUOC, APE, OLE, GEIH"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}:  ")
        run_l.font.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = COLOR_TEXTO
        run_v = p.add_run(value)
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    doc.add_paragraph()

    # Linea decorativa inferior
    p_linea2 = doc.add_paragraph()
    p_linea2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_linea2.add_run("_" * 60)
    run.font.color.rgb = COLOR_PRINCIPAL
    run.font.size = Pt(10)

    # Salto de pagina
    doc.add_page_break()


def _md_a_docx(doc: Document, markdown_text: str):
    """Convierte Markdown basico a elementos DOCX con estilo."""
    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Saltar lineas vacias
        if not line.strip():
            if in_table and table_rows:
                _renderizar_tabla(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GRIS
            i += 1
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="A09088"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Tablas Markdown
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Saltar separadores de tabla (|:---|---|)
            if not re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            texto = line[2:].strip()
            # Procesar formato inline dentro del blockquote
            _procesar_parrafo_formato(doc, texto, italic=True, indent=True)
            i += 1
            continue

        # Bold text marker **text**
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_PRINCIPAL
            i += 1
            continue

        # Listas
        if re.match(r'^\d+\.\s', line):
            texto = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            _procesar_inline_formato(p, texto)
            i += 1
            continue

        if line.strip().startswith('- ') or line.strip().startswith('* '):
            texto = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            _procesar_inline_formato(p, texto)
            i += 1
            continue

        # Parrafo normal
        _procesar_parrafo_formato(doc, line)
        i += 1

    # Procesar tabla pendiente
    if in_table and table_rows:
        _renderizar_tabla(doc, table_rows)


def _procesar_parrafo_formato(doc, texto: str, italic: bool = False, indent: bool = False):
    """Agrega un parrafo procesando formato inline (negritas, italicas, LaTeX)."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    _procesar_inline_formato(p, texto, italic)


def _procesar_inline_formato(p, texto: str, base_italic: bool = False):
    """Procesa formato inline: **bold**, *italic*, $latex$."""
    # Simplificar: reemplazar marcadores LaTeX y Markdown inline
    # Dividir por patrones de formato
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', texto)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('$') and part.endswith('$'):
            run = p.add_run(part[1:-1])
            run.font.italic = True
            run.font.color.rgb = COLOR_SECUNDARIO
        elif part:
            run = p.add_run(part)
            if base_italic:
                run.font.italic = True


def _renderizar_tabla(doc: Document, rows: list):
    """Renderiza una tabla con formato profesional."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                run.font.name = FONT_CUERPO

                # Primera fila = encabezado
                if i == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    # Fondo oscuro para encabezado
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="9B1B30" w:val="clear"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)
                else:
                    run.font.color.rgb = COLOR_TEXTO
                    if i % 2 == 0:
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="F0EAE4" w:val="clear"/>'
                        )
                        cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # Espacio post-tabla


def generar_reporte_docx(
    contenido_markdown: str,
    nbc: str = None,
    depto: str = None,
    fecha: str = None,
    titulo: str = "INFORME DE PERTINENCIA EDUCATIVA"
) -> bytes:
    """
    Genera un documento Word profesional a partir de contenido Markdown.

    Args:
        contenido_markdown: Texto en formato Markdown del informe
        nbc: Nombre del NBC analizado
        depto: Departamento/region
        fecha: Fecha del informe
        titulo: Titulo del documento

    Returns:
        Bytes del archivo .docx listo para descargar
    """
    if fecha is None:
        fecha = datetime.now().strftime("%d de %B de %Y")

    doc = Document()
    _aplicar_estilo_base(doc)
    _agregar_encabezado(doc)
    _agregar_pie_pagina(doc)
    _agregar_marca_agua(doc, "ESTUDIO CONTEXTO")

    # Portada
    _portada(doc, nbc, depto, fecha)

    # Contenido
    _md_a_docx(doc, contenido_markdown)

    # Guardar a bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
